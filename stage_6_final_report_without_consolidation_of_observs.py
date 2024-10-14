#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jun 27 18:13:54 2024

@author: ashutoshgoenka
"""

import pandas as pd
import numpy as np
import streamlit as st
from PIL import Image , ImageOps, ExifTags
import PIL
import os
import zipfile
from zipfile import ZipFile, ZIP_DEFLATED
import pathlib
import shutil
import docx
import docxtpl
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches
import random
from random import randint
from streamlit import session_state
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from streamlit import session_state

st.set_page_config(layout="wide")




state = session_state
if "key" not in state:
    state["key"] = str(randint(1000, 100000000))

if "photo_saved" not in state:
    state["photo_saved"] = False

if "sample_file" not in state:
    state["sample_file"] = False
    
    
if "location_file" not in state:
    state["location_file"] = False

if "page_first_loaded" not in state:
    state["page_first_loaded"] = True
    
if "row_no" not in state:
    state["row_no"] = 1

    


if state["page_first_loaded"] == True:
    
    try:
        shutil.rmtree("images_comp_audit")
    except:
        pass
    
    try:
        os.mkdir("images_comp_audit")
    except:
        pass
    
    state["page_first_loaded"] = False
    
    
def createfile():
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        section.top_margin = Cm(4)
        document.save("sample_output.docx")
        state["photo_saved"] = False
        state["sample_file"] = True




    

def createfile_location():
    if state["location_file"] == False:
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1)
        section.right_margin = Cm(0.75)
        section.bottom_margin = Cm(1)
        document.save("location_output.docx")
        state["location_file"] = True





def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width
        

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'bottom', 'end','insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                    
def set_rows_cant_split(row):
    tr = row._tr
    cantSplits = tr.xpath("./w:trPr/w:cantSplit")
    if cantSplits:
        cantSplit = cantSplits[0]
        cantSplit.set(qn('w:val'), 'true')
        return
    
    
def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def row_position(df_final):
    # df_final= df2
    # st.write(df_final)
    identical_rows_dict = {}
    row_positions  = []
    item_obs = df_final["Element"] + df_final["Observations"]
    item_obs_list = list(item_obs)
    # st.write(item_obs_list)
    for i,j in item_obs.value_counts().items():
        identical_rows_dict[i] = j
    #         print(row_positions[-1])
        start_pos = item_obs_list.index(i) + 1
        end_pos = start_pos + j -1
       
        row_positions.append((start_pos, end_pos))
    return row_positions


def allowDocumentBreak(document):
    """Allow table rows to break across pages."""
    tags = document.element.xpath("//w:tr")
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement("w:cantSplit")  # Create arbitrary tag
#         child.set(qn("w:val"), "0")
        tag.append(child)  # Append in the new tag
        
        
def updateTable_final(df_final, test):
    global df2
    global final_col_width
    doc = docx.Document('sample_output.docx')
    # df_final = df2.copy()
    # df_final.to_excel("data_final_for_table.xlsx")
    t = doc.add_table(df_final.shape[0]+1, df_final.shape[1])
    t.style = 'Table Grid'
    t.allow_autofit = False
    for i in range(len(df_final.columns)-1):
        for cell in t.columns[i].cells:
            cell.width = Cm(float(final_col_width[i]))
    
    
    
    # for cell in t.columns[0].cells:
    #     cell.width = Cm(1.5)
    # for cell in t.columns[1].cells:
    #     cell.width = Cm(3)
    # for cell in t.columns[2].cells:
    #     cell.width = Cm(3.5)
    # for cell in t.columns[3].cells:
    #     cell.width = Cm(4.5)
    # for cell in t.columns[4].cells:
    #     cell.width = Cm(7.75)
    # for cell in t.columns[5].cells:
    #     cell.width = Cm(2)
    # for cell in t.columns[6].cells:
    #     cell.width = Cm(1.75)
    # for cell in t.columns[7].cells:
    #     cell.width = Cm(2)
    
    # add the header rows.
    for j in range(df_final.shape[-1]):
        t.cell(0,j).text = df_final.columns[j]
    
    ctr_temp_1 = 1
    # add the rest of the data frame
    for i in range(df_final.shape[0]):
        for j in range(df_final.shape[-1]):
            if j == 0:
                t.cell(i+1,j).text = str(ctr_temp_1)
                ctr_temp_1 = ctr_temp_1 + 1
            else:
                t.cell(i+1,j).text = str(df_final.values[i,j])
    
    # st.write(df_final)
    row_positions = row_position(df_final)
    # st.write(row_positions)
    
    st_points_list = [i[0] for i in row_positions]
    st_points_list.sort()
    sort_dict = {st_points_list[i]: i for i in range(len(st_points_list))}
    
    
    for row_no in row_positions:
        # st.write(row_no)
        st_pos = row_no[0]
        end_pos = row_no[1]
        ctr_row = 0
        if st_pos != end_pos:
            for i in [0,1,2,3,7]: 
                a = t.cell(st_pos, i)
                temp_content = a.text
    
    #             b = t.cell(2, i)
                c = t.cell(end_pos, i)
            #  # Delete text in cell before merging
            #     delete_paragraph(b.paragraphs[0])
            #     delete_paragraph(c.paragraphs[0])
            #    c.text = ""
            #    A = a.merge(c)
            #    A.text = temp_content
                # if i == 0:
                #     for pt in range(st_pos, end_pos+1):
                #         t.cell(pt, i).text = str(pt)
                #         # ctr_row = ctr_row + 1
            for col in [4,5,6]:
                 for row in range(st_pos,end_pos+1):
                     temp_cell = t.cell(row, col)
             #         if row < end_pos:
             #             set_cell_border(
             #                     temp_cell,
             #             #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
             #                     bottom={"sz": 5, "color": "#E6EDF3", "val": "single"},
             #     #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
             #     #                 end={"sz": 9, "color": "#000000", "val": "single"},
             #                 )
             #         set_cell_border(
             #                 temp_cell,
             #         #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
             # #                 bottom={"sz": 9, "color": "#F0F4F8", "val": "single"},
             # #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
             #                 start={"sz": 6, "color": "#000000", "val": "single"},
             #                 end={"sz": 6, "color": "#000000", "val": "single"},
             #             )
            # add font color to Category
        # else:
        #     a = t.cell(st_pos, 0)
        #     a.text =  str(st_pos)
            
            
        for col in [4,5,6]:
                for cell in t.columns[col].cells:
                    if cell.text == "Alert":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(25,140,25)
                    elif cell.text == "Alarm":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,255)
                    elif cell.text == "Emergency":
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,0,0)
                #     cell.paragraphs[0].paragraph_format.line_spacing = 1.5
                    cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
                    cell.paragraphs[0].paragraph_format.space_after = Cm(0.25)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if col in [5,6]:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        for col in [0,1,2,3,7]:
                for cell in t.columns[col].cells:
                    cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
                    cell.paragraphs[0].paragraph_format.space_after = Cm(0.3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if col in [0,1,7]:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
            # for col in [4,5,6]:
            #     for row in range(st_pos,end_pos+1):
            #         temp_cell = t.cell(row, col)
            #         if row < end_pos:
            #             set_cell_border(
            #                     temp_cell,
            #             #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            #                     bottom={"sz": 5, "color": "#E6EDF3", "val": "single"},
            #     #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
            #     #                 end={"sz": 9, "color": "#000000", "val": "single"},
            #                 )
            #         set_cell_border(
            #                 temp_cell,
            #         #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            # #                 bottom={"sz": 9, "color": "#F0F4F8", "val": "single"},
            # #                 start={"sz": 24, "val": "dashed", "shadow": "true"},
            #                 start={"sz": 6, "color": "#000000", "val": "single"},
            #                 end={"sz": 6, "color": "#000000", "val": "single"},
            #             )
           
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(t.rows[0])
    for row in t.rows:
        set_rows_cant_split(row)
        
        
    
    state["row_no"] = state["row_no"] + len(st_points_list)
    allowDocumentBreak(doc)
    if test == True:
        doc.add_page_break()
    doc.save('./sample_output.docx')   


img_num_dict = {}



def change_state():
    if state["photo_saved"] == True:
        state["photo_saved"] == False
        
def save_image(df,up_files):
    global img_num_dict
    st.write(df)
    img_list = []
    img_num_list = []
    location_dict = []
    segment_list = []
    for idx, rows in df.iterrows():
        if str(rows[10]) != "nan" and str(rows[10])!="-":
            temp_img_list = str(rows[10]).split(",")
            temp_img_list = [t.strip() for t in temp_img_list]
            temp_img_num = str(rows[6]).split("-")
            temp_img_num = [int(t.strip()) for t in temp_img_num]
            if len(temp_img_num)>1:
                temp_img_num = [t for t in range(temp_img_num[0], temp_img_num[1]+1) ]
            if pd.isna(rows[4]):
                temp_loc = "NoAv"
            else:
                temp_loc = rows[4]
            try:
                temp_seg = rows[11]
            except:
                pass
            for m in range(len(temp_img_num)):
                location_dict.append(temp_loc)
                try:
                   segment_list.append(temp_seg) 
                except:
                    pass
                
            img_list = img_list + temp_img_list
            img_num_list = img_num_list + temp_img_num
    # st.write(img_num_list)
    st.write(img_list)
    # st.write(img_num_list)
    
    # st.wire()
    for j in range(len(img_list)):
        
        
        try:
            img_num_dict[img_list[j]] = [img_num_list[j], location_dict[j], 0,0, segment_list[j]]
        except:
            img_num_dict[img_list[j]] = [img_num_list[j], location_dict[j],0,0]
    
    for file in up_files:
        file_name  =file.name
        ext_name =  file_name.split(".")[-1]
        im2 = Image.open(file)
        im2 = ImageOps.expand(im2, border=3)  #Add border to the images
        for img_t in img_list:
            if img_t in file_name:
                loc_name = img_num_dict[img_t][1]
                # st.write(img_t, loc_name)
                
                loc_name_list = loc_name.split(".")
                new_loc_name = ""
                for m in loc_name_list:
                    new_loc_name = new_loc_name + str(m)
                    
                try:
                    new_loc_name = new_loc_name + "_"+img_num_dict[img_t][4]
                except:
                    pass
                new_file_name = "Img "+ str(img_num_dict[img_t][0]) + " @@@" + new_loc_name + "."+ ext_name
                new_file_name = new_file_name.replace("/", "-")
                img_num_dict[img_t][2] = new_file_name
                img_num_dict[img_t][3] = file_name
                # img_num_dict[img_t].append(new_file_name)
                # img_num_dict[img_t].append(file_name)
                
                im2.save("images_comp_audit/"+new_file_name)
    # st.write(img_num_dict)
                
                
def update_test_table():
    createfile()
    state["row_no"] = 1
    global df2
    global up_files
    try:
        segment_list =  list(df2["Segment"])
        segment = []
        for m in segment_list:
            if m not in segment:
                segment.append(m)
        for seg in segment:
            df_seg = df2[df2["Segment"] == seg]
            df_seg = df_seg.copy(deep=True)
            section_report = set(df_seg["Section"])
            state["row_no"] =1
            count_temp = 1
            for i in section_report:
                df3 = df_seg[df_seg["Section"] == i]
                no_of_images = sum(list(df3["No of Images"]))
                df3_temp =  df3.copy(deep=True)
                for idx, row in df3_temp.iterrows():
                    no_of_img_temp =  df3_temp.loc[idx,"No of Images"]
                    if no_of_img_temp>1:
                        start_val_temp =  count_temp
                        end_val_temp = count_temp + no_of_img_temp - 1
                        img_val = "00"+str(start_val_temp) + " - " + "00"+str(end_val_temp)
                        df3_temp.loc[idx,"Image No."] = img_val 
                        count_temp = count_temp + no_of_img_temp
                    else:
                        start_val_temp =  count_temp
                        img_val = "00"+str(start_val_temp)
                        df3_temp.loc[idx,"Image No."] = img_val 
                        count_temp = count_temp + no_of_img_temp
                df3 = df3_temp    
                
                # st.write("no of image" + str(no_of_images))
                df3 = df3.drop(["No of Images","Segment"], axis=1)
                df4 = df3.copy(deep=True)
                df3 = df3.iloc[:, :-2]
                # st.write(df3)
                updateTable_final(df3, True)
                # updateImage(df4, no_of_images, up_files)  
            
            
            
            
    except:
        section_report = set(df2["Section"])
        state["row_no"] =1
        count_temp = 1
        for i in section_report:
            df3 = df2[df2["Section"] == i]
            no_of_images = sum(list(df3["No of Images"]))
            df3_temp =  df3.copy(deep=True)
            for idx, row in df3_temp.iterrows():
                no_of_img_temp =  df3_temp.loc[idx,"No of Images"]
                if no_of_img_temp>1:
                    start_val_temp =  count_temp
                    end_val_temp = count_temp + no_of_img_temp - 1
                    img_val = "00"+str(start_val_temp) + " - " + "00"+str(end_val_temp)
                    df3_temp.loc[idx,"Image No."] = img_val 
                    count_temp = count_temp + no_of_img_temp
                else:
                    start_val_temp =  count_temp
                    img_val = "00"+str(start_val_temp)
                    df3_temp.loc[idx,"Image No."] = img_val 
                    count_temp = count_temp + no_of_img_temp
            df3 = df3_temp    
            
            # st.write("no of image" + str(no_of_images))
            df3 = df3.drop(["No of Images", "Segment"], axis=1)
            df4 = df3.copy(deep=True)
            df3 = df3.iloc[:, :-2]
            # st.write(df3)
            updateTable_final(df3, True)
            # updateImage(df4, no_of_images, up_files)  
    
    
         
def updateTable_new():
    state["row_no"] = 1
    createfile()
    updateWordDoc()
        
 
        
def updateWordDoc():
    global df2
    global up_files
    try:
        segment_list =  list(df2["Segment"])
        segment = []
        for m in segment_list:
            if m not in segment:
                segment.append(m)
        for seg in segment:
            df_seg = df2[df2["Segment"] == seg]
            df_seg = df_seg.copy(deep=True)
            section_report = set(df_seg["Section"])
            state["row_no"] =1
            count_temp = 1
            for i in section_report:
                df3 = df_seg[df_seg["Section"] == i]
                no_of_images = sum(list(df3["No of Images"]))
                df3_temp =  df3.copy(deep=True)
                for idx, row in df3_temp.iterrows():
                    no_of_img_temp =  df3_temp.loc[idx,"No of Images"]
                    if no_of_img_temp>1:
                        start_val_temp =  count_temp
                        end_val_temp = count_temp + no_of_img_temp - 1
                        img_val = "00"+str(start_val_temp) + " - " + "00"+str(end_val_temp)
                        df3_temp.loc[idx,"Image No."] = img_val 
                        count_temp = count_temp + no_of_img_temp
                    else:
                        start_val_temp =  count_temp
                        img_val = "00"+str(start_val_temp)
                        df3_temp.loc[idx,"Image No."] = img_val 
                        count_temp = count_temp + no_of_img_temp
                df3 = df3_temp    
                
                # st.write("no of image" + str(no_of_images))
                df3 = df3.drop(["No of Images","Segment"], axis=1)
                df4 = df3.copy(deep=True)
                df3 = df3.iloc[:, :-2]
                # st.write(df3)
                updateTable_final(df3, True)
                updateImage(df4, no_of_images, up_files)  
    
    
    
    
    except:
        section_report = set(df2["Section"])
        count_temp = 1
        for i in section_report:
            df3 = df2[df2["Section"] == i]
            no_of_images = sum(list(df3["No of Images"]))
            # st.write("no of image" + str(no_of_images))
            df3_temp =  df3.copy(deep=True)
            for idx, row in df3_temp.iterrows():
                no_of_img_temp =  df3_temp.loc[idx,"No of Images"]
                if no_of_img_temp>1:
                    start_val_temp =  count_temp
                    end_val_temp = count_temp + no_of_img_temp - 1
                    img_val = "00"+str(start_val_temp) + " - " + "00"+str(end_val_temp)
                    df3_temp.loc[idx,"Image No."] = img_val 
                    count_temp = count_temp + no_of_img_temp
                else:
                    start_val_temp =  count_temp
                    img_val = "00"+str(start_val_temp)
                    df3_temp.loc[idx,"Image No."] = img_val 
                    count_temp = count_temp + no_of_img_temp
                    
            df3 = df3_temp       
            df3 = df3.drop(["No of Images"], axis=1)
            
            
            
            
            df4 = df3.copy(deep=True)
            df3 = df3.iloc[:, :-2]
            # st.write(df3)
            updateTable_final(df3, False)
            updateImage(df4, no_of_images, up_files)


def updateImage(df3, no_of_images, up_files):
    doc = docx.Document('sample_output.docx')
    doc.add_heading("Images", 2)
    no_of_rows = int(((no_of_images-1)//3+1)*2)   ## 3 is no of columns
    table = doc.add_table(rows = no_of_rows , cols = 3)
    # test
    for col in table.columns:
        for cell in col.cells:
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            
    ctr_temp = 0
    # for row in table.rows:
    #     if ctr_temp %2 ==0:
    #         row.height = Cm(8.5)
    #     else:
    #         row.height = Cm(1.25)
    #     ctr_temp += 1
    counter = 0
    counter_cols = 0
    image_list = list(df3["Image No."])
    location_list = list(df3["Location"])
    image_location_dict = {}
    img_list = []
    for j in range(len(image_list)):
        temp_str_list = image_list[j].split("-")
        temp_loc = location_list[j]
        if len(temp_str_list) == 1:
            temp_start_pt =  int(temp_str_list[0])
            image_location_dict[temp_start_pt] =  temp_loc
            img_list.append(temp_start_pt)
        else:
            temp_start_pt = int(temp_str_list[0])
            temp_end_pt = int(temp_str_list[1])
            for k in range(temp_start_pt, temp_end_pt+1):
                image_location_dict[k] = temp_loc
                img_list.append(k)
        

    image_number_df_list = list(df3["Image Number"])
    # st.write("Test")
    # st.write(image_number_df_list)
    image_number_list = []
    for i in image_number_df_list:
        # st.write()
        if str(i) != "nan" and str(i)!="-":
            temp_list_img = (str(i).split(","))
            temp_list_img = [t.strip() for t in temp_list_img]  
            image_number_list = image_number_list + temp_list_img
        
    t_list = []
    # st.write(image_number_list)    
    for i in range(len(image_number_list)):
        # st.write(image_number_list)
        img_no = image_number_list[i]
        t_list.append(img_no)
        try:
            image_file_t = img_num_dict[img_no]
        except:
            st.write(t_list)
            raise ValueError
        # image_file_name = image_file_t[2]
        image_file_name = img_num_dict[img_no][2]
        file_temp1 = 'images_comp_audit/'+image_file_name
        # st.write("open ", file_temp1 )
        im_temp = Image.open(file_temp1)
        im_width, im_height = im_temp.size
        
        # for file in up_files:
        #     if image_name == file.name:
        #         im_temp = Image.open(file)
        row_no = (i//3) *2
        col_no = int(i - (row_no*3/2))
        cell = table.rows[row_no].cells[col_no]
        cell._element.clear_content()
        # st.write(image_file_name)
        # st.write(im_width, im_height)
        if im_width>=im_height:
            picture = cell.add_paragraph().add_run().add_picture('images_comp_audit/'+image_file_name, width=Cm(6.5))
        else:
            picture = cell.add_paragraph().add_run().add_picture('images_comp_audit/'+image_file_name, height=Cm(6.5))
        cell = table.rows[row_no+1].cells[col_no]
        # cell = table.rows[counter+1].cells[counter_cols]
        # st.write(row_no, col_no)
        try:
            name = "Img " + str(img_list[i]) +" at " + image_location_dict[img_list[i]]
        except:
            name = "Img " + str(img_list[i])
        cell.text = name
        if col_no<2:
            counter_cols = counter_cols + 1
        else:
            # table.add_row()
            counter_cols = 0
            counter = counter+2
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in table.columns:
        for cell in col.cells:
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.save('./sample_output.docx')  



def updateLocationTable():
    global df2
    df_loc = 1
    



# Function for updating the location word file
def updateLocationWordDoc():
    try:
        createfile_location()
    except:
        pass
    
    global df_loc_final
    global up_files
    section_report = set(df_loc_final["Main Section"])
    
    for i in section_report:
        
        df3 = df_loc_final[df_loc_final["Main Section"] == i]
        no_of_images = sum(list(df3["No of Images"]))
        # st.write("no of image " + str(no_of_images))
        df3 = df3.drop(["No of Images"], axis=1)
        location_temp =  list(df3["Location"])[0]
        df4 = df3.copy(deep=True)
        df4 = df4.iloc[:, :-4]
        df4 = df4.drop(["Location"], axis=1)
        df4 =df4.reset_index(drop =True)
        for idx,rows in df4.iterrows():
            df4.loc[idx, "Sl"] = idx+1
        # st.write(df4)
        upadateLocationTable(df4, location_temp) 
        updateImageLocation(df3, no_of_images, up_files)

def upadateLocationTable(df_final_2, loc_temp):
    global df_loc_final
    global final_col_width
    doc = docx.Document('location_output.docx')
    # df_final_2 = df_loc_final.copy(deep = True)
    # df_final_2 = df_final_2
    # df_final.to_excel("data_final_for_table.xlsx")
    
    doc.add_heading("Location: " + loc_temp, 2)
    t = doc.add_table(df_final_2.shape[0]+1, df_final_2.shape[1])
    t.style = 'Table Grid'
    t.allow_autofit = False
    for i in range(len(df_final_2.columns)-1):
        for cell in t.columns[0].cells:
            cell.width = Cm(1.5)
        for cell in t.columns[1].cells:
            cell.width = Cm(3)
        for cell in t.columns[2].cells:
            cell.width = Cm(7.5)
        for cell in t.columns[3].cells:
            cell.width = Cm(7.5)
        for cell in t.columns[4].cells:
            cell.width = Cm(2.5)
        for cell in t.columns[5].cells:
            cell.width = Cm(2)
        for cell in t.columns[6].cells:
            cell.width = Cm(2)
        
          # add the header rows.
        for j in range(df_final_2.shape[-1]):
              t.cell(0,j).text = df_final_2.columns[j]
          
          # add the rest of the data frame
        for i in range(df_final_2.shape[0]):
              for j in range(df_final_2.shape[-1]):
                  t.cell(i+1,j).text = str(df_final_2.values[i,j])
          
        for col in [0,1,2,3,4,5,6]:
                for cell in t.columns[col].cells:
                    cell.paragraphs[0].paragraph_format.space_before = Cm(0.25)
                    cell.paragraphs[0].paragraph_format.space_after = Cm(0.3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if col in [0,1,4,5,6]:
                        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        row_positions = row_position(df_final_2)
    # doc.add_page_break()
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(t.rows[0])
    for row in t.rows:
        set_rows_cant_split(row)
    
    
    doc.save('./location_output.docx')  
    


def updateImageLocation(df3, no_of_images, up_files):
    doc = docx.Document('location_output.docx')
    doc.add_heading("Images", 2)
    no_of_rows = int(((no_of_images-1)//3+1)*2)   ## 3 is no of columns
    table = doc.add_table(rows = no_of_rows , cols = 3)
    ctr_temp = 0
    # for row in table.rows:
    #     if ctr_temp %2 ==0:
    #         row.height = Cm(8.5)
    #     else:
    #         row.height = Cm(1.25)
    #     ctr_temp += 1
    counter = 0
    counter_cols = 0
    image_list = list(df3["Image No."])
    location_list = list(df3["Location"])
    obs_list = list(df3["Observations"])
    image_location_dict = {}
    image_obs_dict = {}
    img_list = []
    for j in range(len(image_list)):
        temp_str_list = image_list[j].split("-")
        temp_loc = location_list[j]
        temp_obs = obs_list[j]
        if len(temp_str_list) == 1:
            temp_start_pt =  int(temp_str_list[0])
            image_location_dict[temp_start_pt] =  temp_loc
            image_obs_dict[temp_start_pt] = temp_obs
            img_list.append(temp_start_pt)
        else:
            temp_start_pt = int(temp_str_list[0])
            temp_end_pt = int(temp_str_list[1])
            for k in range(temp_start_pt, temp_end_pt+1):
                image_location_dict[k] = temp_loc
                image_obs_dict[k] = temp_obs
                img_list.append(k)
        

    image_number_df_list = list(df3["Image Number"])
    # st.write("Test")
    # st.write(image_number_df_list)
    image_number_list = []
    for i in image_number_df_list:
        # st.write()
        temp_list_img = (str(i).split(","))
        temp_list_img = [t.strip() for t in temp_list_img]  
        image_number_list = image_number_list + temp_list_img
        

    # st.write(image_number_list)    
    for i in range(len(image_number_list)):
        # st.write(image_number_list)
        img_no = image_number_list[i]
        image_file_t = img_num_dict[img_no]
        image_file_name = image_file_t[2]
        # image_file_name = img_num_dict[img_no][2]
        file_temp1 = 'images_comp_audit/'+image_file_name
        im_temp = Image.open(file_temp1)
        im_width, im_height = im_temp.size
        
        # for file in up_files:
        #     if image_name == file.name:
        #         im_temp = Image.open(file)
        row_no = (i//3) *2
        col_no = int(i - (row_no*3/2))
        cell = table.rows[row_no].cells[col_no]
        cell._element.clear_content()
        # st.write(image_file_name)
        # st.write(im_width, im_height)
        if im_width>=im_height:
            picture = cell.add_paragraph().add_run().add_picture('images_comp_audit/'+image_file_name, width=Cm(8.25))
        else:
            picture = cell.add_paragraph().add_run().add_picture('images_comp_audit/'+image_file_name, height=Cm(8.25))
        cell = table.rows[row_no+1].cells[col_no]
        # cell = table.rows[counter+1].cells[counter_cols]
        # st.write(row_no, col_no)
        name = "Img " + str(img_list[i]) +" - " + image_obs_dict[img_list[i]]
        cell.text = name
        if col_no<2:
            counter_cols = counter_cols + 1
        else:
            # table.add_row()
            counter_cols = 0
            counter = counter+2
    doc.add_page_break()
    doc.save('./location_output.docx')  

st.title("Stage 6 - Final Report Preperation without Consolidation of Observation")
### Main Page Starts from here
st.write("Upload Optimized Observation")
obs_file = st.file_uploader("Upload Observation Excel File", type=['csv','xlsx'],accept_multiple_files=False,key="fileUploader")
if obs_file is not None:
    # uplaod remedy file
    df_rem = pd.read_excel("remedy_excel.xlsx")
    remedy_dict = {}
    for idx, val in df_rem.iterrows():
    #     print(val)
        temp_key = val["Observations"] + val["Severity"]
        remedy_dict[temp_key] = (val["Remedy"], val["Category"])
        
    
    df = pd.read_excel(obs_file)
    df.drop(df.columns[[0]], axis=1, inplace=True)
    df = df.dropna(thresh=5)
    image_master_list_obs = []
    image_row_dict = {}
    for idx, rows in df.iterrows():
        temp_img_list = str(rows[10]).split(",")
        temp_img_list = [t.strip() for t in temp_img_list]
        image_row_dict[rows[0]] = temp_img_list
        image_master_list_obs = image_master_list_obs + temp_img_list
        df.loc[idx, "Image Number"] = str(rows[10])
    st.title("The Uploaded Data:")
    st.write(df)
    # st.write(image_master_list_obs)
    # st.write(image_row_dict)
    df2 =df.copy(deep=True)
    # df2= df2.drop(["No of Images"], axis=1)
    # st.write(df2)
    section_report = set(df2["Section"])
    # save_image()
    
    #Upload the images
    
    st.write("Table Column Width in cm (Avoid Changing it as it was used for optimization)")
    no_of_cols = len(df2.columns)
    default_co_width = [1.5, 3.0, 3.5, 4.5, 7.75, 2, 1.75, 2]
    col_list = st.columns(8)  #hard coded the number of columns 
    final_col_width= [i for i in default_co_width]
    total_width = sum([float(i) for i in final_col_width])
    for i in range(8):   #hard coded the number of columns 
        with col_list[i]:
            final_col_width[i] = st.text_input(df2.columns[i],
                          default_co_width[i],
                          key="col_data_"+str(i),
    )
    if final_col_width[0]:
        total_width = sum([float(i) for i in final_col_width])
    st.write("Total Column Width: "+ str(total_width))
    # st.write(final_col_width)
    
    if state["sample_file"] == False:
        createfile()
    
    
    try:
        with open("sample_output.docx", "rb") as fp:
        
            btn_1 = st.button(
                    label="Create New Word File and Update Data in Test Table",
                    on_click=update_test_table      
                )
    except:
        pass
    
    try:
        with open("sample_output.docx", "rb") as fp:
        
            btn_1 = st.download_button(
                    label="Download Test Word File",
                    data=fp,
                    file_name="sample_output.docx",
                    mime="docx"
                    )
    except:
        pass
    
    
    
    
    up_files = st.file_uploader("Upload Image Files", type = ["png", "jpeg", "jpg"] ,accept_multiple_files=True, key=state["key"], on_change=change_state)
    all_file_found = False
    if len(up_files)>=0:
        file_name_list = []
        missing_img = []
        found_img = []
        for file in up_files:
            file_name_list.append(file.name)
        for img in image_master_list_obs:
            for file_temp in file_name_list:
                if img in file_temp:
                    found_img.append(img)
            if img not in found_img and img != "nan":
                missing_img.append(img)
        
        if len(missing_img)>0:
            st.write("The following images are missing: ", str(missing_img))
        elif len(up_files)>0 and len(missing_img)==0 and state["photo_saved"] == False:
            save_image(df,up_files)
            state["photo_saved"] == True
            
        
            
    # st.write("Test 2")
    # st.write(section_report)
    
    
        
        if len(missing_img)==0:
            if state["sample_file"] == False:
                createfile()
            
            
            try:
                with open("sample_output.docx", "rb") as fp:
                
                    btn_1 = st.button(
                            label="Create New Word File and Update Data",
                            on_click=updateTable_new       
                        )
            except:
                pass
            
            
            
            
            # try:
            #     with open("sample_output.docx", "rb") as fp:
                
            #         btn_1 = st.button(
            #                 label="Update Word File",
            #                 on_click=updateWordDoc,       
            #             )
            # except:
            #     pass
            
            
            try:
                with open("sample_output.docx", "rb") as fp:
                
                    btn_1 = st.download_button(
                            label="Download Word File",
                            data=fp,
                            file_name="sample_output.docx",
                            mime="docx"
                            )
            except:
                pass
            
           
            
           
            st.write("Generate Table by Location")
            obs_file_2 = st.file_uploader("Upload New Location File (Optional)", type=['csv','xlsx'],accept_multiple_files=False,key="fileUploaderlocation")
            if obs_file_2 is None:
                df_loc_final = df2.copy(deep = True)
            else:
                df_loc_final= pd.read_excel(obs_file)
            
            df_loc_final = df_loc_final.sort_values(by=["Location"])
            # st.write(df_loc_final)
            location_list = list(set(df_loc_final["Location"]))
            location_order = [i+1 for i in range(len(location_list))]
            loc_order_dict = {"Location" : location_list,
                              "Location Order": 0}
            df_location = pd.DataFrame(loc_order_dict)
            df_location = df_location.sort_values(by=["Location"])
            df_location["Location Order"] = location_order
            # st.write(df_location)
            st.write("Please select the order of the location for the report")
            df_location = st.data_editor(
                    df_location,
                    column_config={
                        "Location Order": st.column_config.NumberColumn(
                            "Order of Table in Word",
                            help="The order in which the table will be added to word",
                            min_value=1,
                            max_value=len(location_list),
                            step=1,
                            format="%d",
                                )
                            },
                            hide_index=True,
                            )
            order_list =  list(df_location["Location Order"])
            all_values_selected = True
            for i in range(1, len(location_list)+1):
                if i not in order_list:
                    all_values_selected = False
            
            # st.write(order_list)
            # st.write(all_values_selected)
            loc_order = {}
            if all_values_selected == False:
                st.write("Please mention the proper order in the above table")
                
            else:
                for idx, rows in df_location.iterrows():
                    loc_order[rows[0]] = rows[1]
                    
                    
                    
                    
                df_loc_final_temp = df_loc_final.loc[:,["Location", "Element", "Observations"]]
                df_loc_final_temp["Location Order"] = 0
                for idx, rows in df_loc_final_temp.iterrows():
                    order_val = loc_order[rows[0]]
                    df_loc_final_temp.loc[idx, "Location Order"] =  order_val
                
                
                df_loc_final_temp = df_loc_final_temp.sort_values(by =["Location Order"])
                df_loc_final_temp["Section Order"] = 0
                sec_order = 1
                ctr_order = 1
                for idx, rows in df_loc_final_temp.iterrows():
                    new_sec_order = rows[3]
                    if new_sec_order == sec_order:
                        df_loc_final_temp.loc[idx, "Section Order"] = ctr_order
                        ctr_order = ctr_order + 1
                    else:
                        ctr_order = 1
                        sec_order= rows[3]
                        df_loc_final_temp.loc[idx, "Section Order"] = ctr_order
                        ctr_order = ctr_order + 1
                        
                    
                st.write("Realigned Rows as per the order provided above")
                # st.write(df_loc_final_temp)
                df_loc_final_temp = st.data_editor(
                        df_loc_final_temp,
                        column_config={
                            "Section Order": st.column_config.NumberColumn(
                                "Section Order",
                                help="The order in which the table will be added to word",
                                min_value=0,
                                max_value=100,
                                step=1,
                                format="%d",
                                    )
                                },
                                hide_index=True,
                                )
                
                
                temp_master_dict = {}
                for idx, rows in df_loc_final_temp.iterrows():
                    temp_str =  rows[0] + rows[1]+ rows[2]
                    temp_master_dict[temp_str] = (rows[3], rows[4])
                # st.write(temp_master_dict)
                
                df_loc_final["Main Section"] = 0
                df_loc_final["Sub Section"] = 0
                for idx , rows in df_loc_final.iterrows():
                    loc_temp = rows[4]
                    element_temp = rows[1]
                    obs_temp = rows[2]
                    temp_str =  loc_temp+element_temp+obs_temp
                    temp_result = temp_master_dict[temp_str]
                    df_loc_final.loc[idx, "Main Section"] = temp_result[0]
                    df_loc_final.loc[idx, "Sub Section"] = temp_result[1]
                # st.write(df_loc_final)
                sub_sec_order_correct = True
                df_temp = df_loc_final.filter(["Main Section", "Sub Section"])
                main_sec_list = set(df_temp["Main Section"])
                for sec in  main_sec_list:
                    df_temp2  =   df_temp[df_temp["Main Section"] == sec]
                    sub_sec = list(df_temp2["Sub Section"])
                    for sub_ord_val in range(1, len(sub_sec)+1):
                        if sub_ord_val not in sub_sec:
                            sub_sec_order_correct = False
                            break
                        else:
                            pos = sub_sec.index(sub_ord_val)
                            sub_sec.pop(pos)
                    if sub_sec_order_correct == False:
                        break
                
                
                # st.write(sub_sec_order_correct)
                if sub_sec_order_correct== True:
                    # st.write(df_temp)
                    df_loc_final= df_loc_final.sort_values(by=["Main Section", "Sub Section"])
                    df_loc_final = df_loc_final.reset_index(drop=True)
                    # st.write("df loc final")
                    # st.write(df_loc_final)
                    img_list =list( df_loc_final["No of Images"])
                    ctr_temp =1
                    for j in range(len(img_list)):
                        no_of_img = img_list[j]
                        if no_of_img>1:
                            df_loc_final.loc[j, "Image No."] = "00"+str(ctr_temp) + " - 00"+str(ctr_temp+no_of_img-1)
                            ctr_temp = ctr_temp+no_of_img
                        else:
                            df_loc_final.loc[j, "Image No."] = "00"+str(ctr_temp)
                            ctr_temp =  ctr_temp + 1
                    
                    # st.write(df_loc_final)   
                    # updateLocationWordDoc()
                    
                    if len(missing_img)==0:
                        if state["location_file"] == False:
                            createfile_location()
                        
                        
                        try:
                            with open("location_output.docx", "rb") as fp:
                            
                                btn_1 = st.button(
                                        label="Create New Word File for Location and Update Data",
                                        on_click=updateLocationWordDoc       
                                    )
                        except:
                            pass
                    
                    
                    
                    
                     
                        try:
                            with open("location_output.docx", "rb") as fp:
                            
                                btn_1 = st.download_button(
                                        label="Download Location Word File",
                                        data=fp,
                                        file_name="location_output",
                                        mime="docx"
                                        )
                        except:
                            pass
                    
                
            # try:
            #     with open("location_table_file.docx", "rb") as fp:
                
            #         btn_1 = st.button(
            #                 label="Update Location Table File",
            #                 on_click=updateLocationTable,       
            #             )
            # except:
            #     pass
    
    # for i in section_report:
    #     df3 = df2[df2["Section"] == i]
    #     no_of_images = sum(list(df3["No of Images"]))
    #     st.write("no of image" + str(no_of_images))
    #     df3 = df3.drop(["No of Images"], axis=1)
    #     df4 = df3.copy(deep=True)
    #     df3 = df3.iloc[:, :-2]
    #     st.write(df3)
    #     updateTable_final(df3)
    #     updateImage(df4, no_of_images, up_files)
        
    # st.write(df3)
   
        
    

